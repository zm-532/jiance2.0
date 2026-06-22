"""Word 报告生成服务：基于模板填充检测数据。

使用 python-docx 打开模板，替换表格中的检测数据，生成最终报告。
"""
import logging
import uuid
from datetime import datetime
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.photo import Photo
from app.services.config_cache import get_image_description

logger = logging.getLogger(__name__)

# 模板目录
TEMPLATE_DIR = Path(__file__).resolve().parent.parent.parent.parent / "docs" / "检测报告模板"

# 模板映射
TEMPLATE_MAP = {
    "R001": "检测报告模版-1.docx",
    "R002": "检测报告模版-2.docx",
    "R003": "检测报告模版-3.docx",
}


async def generate_report(
    db: AsyncSession,
    template_id: str,
    entrust_no: str | None = None,
    sample_name: str | None = None,
) -> BytesIO:
    """生成 Word 检测报告。

    Args:
        db: 数据库会话
        template_id: 模板 ID（如 R001）
        entrust_no: 委托编号（可选过滤）
        sample_name: 样品名称（可选过滤）

    Returns:
        BytesIO: 生成的 Word 文件流
    """
    template_file = TEMPLATE_MAP.get(template_id)
    if not template_file:
        raise ValueError(f"未知模板 ID: {template_id}")

    template_path = TEMPLATE_DIR / template_file
    if not template_path.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    # 查询纳入报告的照片（按 group_id 聚合）
    stmt = (
        select(Photo)
        .where(Photo.include_in_report == True)
        .where(Photo.group_id.isnot(None))
        .order_by(Photo.group_id, Photo.test_item)
    )
    if entrust_no:
        stmt = stmt.where(Photo.entrust_no == entrust_no)
    if sample_name:
        stmt = stmt.where(Photo.sample_name == sample_name)

    result = await db.execute(stmt)
    photos = result.scalars().all()

    if not photos:
        raise ValueError("没有纳入报告的检测数据，请先在「照片OCR」中上传并确认数据")

    # 按 group_id 分组
    groups: dict[str, list[Photo]] = {}
    for p in photos:
        groups.setdefault(p.group_id, []).append(p)

    # 打开模板
    doc = Document(str(template_path))

    # 填充检测结果表（通常是第二个表格，索引 1）
    _fill_results_table(doc, groups)

    # 填充样品信息表（通常是第一个表格，索引 0）
    _fill_sample_info_table(doc, photos, entrust_no)

    # 保存到内存
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    logger.info(f"报告生成完成: 模板={template_id}, 组数={len(groups)}, 检测项={len(photos)}")
    return output


def _fill_results_table(doc: Document, groups: dict[str, list[Photo]]) -> None:
    """填充检测结果表。

    模板中的检测结果表通常有列：序号 | 检测项目 | 检测项目(子) | 标准要求 | 检测值 | 判定
    """
    if len(doc.tables) < 2:
        logger.warning("模板中未找到足够的表格，跳过检测结果填充")
        return

    table = doc.tables[1]  # 第二个表格是检测结果表

    # 收集所有检测项（按组排列）
    all_items = []
    seq = 1
    for group_id, group_photos in groups.items():
        for p in group_photos:
            all_items.append({
                "seq": seq,
                "test_item": p.test_item or "-",
                "sub_item": p.sub_item or "",
                "standard_requirement": p.standard_requirement or "-",
                "recognized_value": p.recognized_value or "-",
                "judgment": p.judgment or "待判定",
            })
            seq += 1

    # 保留表头行（第 0 行），清空其余行后重新填充
    if len(table.rows) <= 1:
        logger.warning("检测结果表无数据行")
        return

    # 删除现有数据行（保留表头）
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    # 添加新数据行
    for item in all_items:
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(item["seq"])
        cells[1].text = item["test_item"]
        cells[2].text = item["sub_item"]
        cells[3].text = item["standard_requirement"]
        cells[4].text = item["recognized_value"]
        cells[5].text = item["judgment"]

        # 设置字号
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _fill_sample_info_table(doc: Document, photos: list[Photo], entrust_no: str | None) -> None:
    """填充样品信息表（第一个表格）。

    尝试找到关键字段并填充，跳过无法匹配的字段。
    """
    if len(doc.tables) < 1:
        return

    table = doc.tables[0]
    first = photos[0] if photos else None
    if not first:
        return

    # 遍历表格，找到标签行并填充对应的值
    fill_map = {
        "样品名称": first.sample_name or "",
        "委托单位": "",  # 需要外部数据
        "样品编号": first.entrust_no or entrust_no or "",
        "委托编号": first.entrust_no or entrust_no or "",
    }

    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = cells[0].text.strip()
        for key, value in fill_map.items():
            if key in label and value:
                cells[1].text = value
                break
